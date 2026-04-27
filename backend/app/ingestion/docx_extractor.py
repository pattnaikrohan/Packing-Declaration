"""
DOCX extractor using python-docx.
Reads paragraphs and table cells.

Checkbox detection has two layers:
  1. Primary  — XML w:checkBox elements (AcroForm-style interactive DOCX)
  2. Fallback — Unicode checkbox characters (☑ U+2611, ☒ U+2612 = ticked; ☐ U+2610 = unticked)
     plus square-bracket OCR surrogates often used in older Word templates.
"""
import re
from lxml import etree
from docx import Document
from io import BytesIO

from app.ingestion.schema import PackingDeclaration
from app.ingestion import extractors_common as ec

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Characters that Tesseract / copy-paste renders as an EMPTY (unticked) checkbox
EMPTY_BOX_CHARS = re.compile(r'^[\s\[\]|LlCcJj_\{\}°\u25a1\u2610]{0,8}$')

# Unicode ticked characters
TICKED_CHARS = {"☑", "☒", "✓", "✔", "✗", "✘"}
UNTICKED_CHARS = {"☐"}


def _is_checked_xml(element) -> bool:
    """Check w:checkBox/w:checked val attribute (interactive DOCX forms)."""
    for cb in element.iter(f"{{{W_NS}}}checkBox"):
        checked_el = cb.find(f"{{{W_NS}}}checked")
        if checked_el is not None:
            val = checked_el.get(f"{{{W_NS}}}val", "1")
            return val not in ("0", "false")
        return True
    return False


def _paragraph_is_ticked_unicode(para) -> bool | None:
    """
    Returns True if paragraph contains a ticked Unicode/Word checkbox char,
    False if it contains an unticked one, None if no checkbox chars found.
    """
    text = para.text
    for ch in TICKED_CHARS:
        if ch in text:
            return True
    for ch in UNTICKED_CHARS:
        if ch in text:
            return False
    # Square-bracket surrogate heuristic:
    # Patterns like [X], [✓], [/] = ticked; [ ] or [] = unticked
    m = re.search(r'\[([^\]]{0,5})\]', text)
    if m:
        inner = m.group(1).strip()
        if inner in ('', ' '):
            return False
        if re.search(r'[xX✓✔/\\vV]', inner):
            return True
        return False
    return None


def _extract_checkboxes(doc) -> dict:
    result = {"q1": "BLANK", "q2": "BLANK", "q3": "BLANK", "q4": "ABSENT"}

    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        xml = para._element
        has_xml_checkbox = bool(xml.find(f".//{{{W_NS}}}checkBox"))
        unicode_ticked = _paragraph_is_ticked_unicode(para)

        # Determine ticked state
        if has_xml_checkbox:
            ticked = _is_checked_xml(xml)
        elif unicode_ticked is not None:
            ticked = unicode_ticked
        else:
            continue  # no checkbox in this paragraph

        text = para.text.upper()

        if "Q1" in text or "UNACCEPTABLE" in text or "PROHIBITED" in text:
            result["q1"] = "YES" if ticked else "NO"
        elif "TIMBER" in text:
            result["q2"] = "YES_TIMBER" if ticked else "NO"
        elif "BAMBOO" in text:
            result["q2"] = "YES_BAMBOO" if ticked else result["q2"]
        elif "Q2" in text:
            if ticked and result["q2"] == "BLANK":
                result["q2"] = "YES_TIMBER"
            elif not ticked and result["q2"] == "BLANK":
                result["q2"] = "NO"
        elif "ISPM" in text:
            result["q3"] = "ISPM15" if ticked else result["q3"]
        elif "DAFF" in text and "CERTIF" in text:
            result["q3"] = "DAFF_CERTIFIED" if ticked else result["q3"]
        elif "NOT TREATED" in text or "UNTREATED" in text:
            result["q3"] = "NOT_TREATED" if ticked else result["q3"]
        elif "NOT APPLICABLE" in text or "N/A" in text:
            result["q3"] = "NOT_APPLICABLE" if ticked else result["q3"]
        elif "Q3" in text or "TREATMENT" in text:
            if ticked and result["q3"] == "BLANK":
                result["q3"] = "ISPM15"
        elif "Q4" in text or "CLEANLINESS" in text:
            result["q4"] = "PRESENT" if ticked else "ABSENT"

    return result


def extract(file_bytes: bytes) -> PackingDeclaration:
    doc = Document(BytesIO(file_bytes))

    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    full_text = "\n".join(parts)
    checkboxes = _extract_checkboxes(doc)

    addr, is_po = ec.extract_address(full_text)
    consignment, _ = ec.extract_consignment_link(full_text)
    date_str, date_valid = ec.extract_date(full_text)
    signed, sig_type = ec.detect_signature(full_text)
    alterations, endorsed = ec.detect_alterations(full_text)

    return PackingDeclaration(
        declaration_type=ec.detect_declaration_type(full_text),
        issuer_company=ec.extract_company(full_text),
        issuer_address=addr,
        issuer_address_is_po_box=is_po,
        vessel_name=ec.extract_vessel(full_text),
        voyage_number=ec.extract_voyage(full_text),
        consignment_ref=consignment,
        exporter=ec.extract_party(full_text, "Exporter"),
        importer=ec.extract_party(full_text, "Importer"),
        date_issued=date_str,
        date_valid=date_valid,
        signed=signed,
        signature_type=sig_type,
        printed_name=ec.extract_printed_name(full_text),
        letterhead_present=ec.detect_letterhead(full_text),
        q1_unacceptable_material=checkboxes["q1"],
        q2_timber_bamboo=checkboxes["q2"],
        q3_treatment=checkboxes["q3"],
        q4_cleanliness=checkboxes["q4"],
        alterations_present=alterations,
        alterations_endorsed=endorsed,
        extraction_method="docx",
        ocr_confidence=1.0,
    )
